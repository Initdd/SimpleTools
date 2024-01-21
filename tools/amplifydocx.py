
from docx import Document
from docx.shared import Pt

def resize_font(docx_src, docx_dest, resize_factor):
    # Load the DOCX document
    doc = Document(docx_src)
    last_font_size = doc.paragraphs[0].runs[0].font.size
    if last_font_size is None:
        last_font_size = Pt(11)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.size is not None:
                current_font_size = run.font.size
                last_font_size = current_font_size
            else:
                current_font_size = last_font_size
            amplified_font_size = current_font_size * resize_factor
            run.font.size = int(amplified_font_size)

    # Save the modified document
    output_path = f'{docx_dest}.docx'
    doc.save(output_path)
    return f"Resized document saved to: {output_path}"

if __name__ == "__main__":
    # Input: DOCX file path and resizing factor
    docx_path = input("Enter the DOCX file path: ")
    resize_factor = float(input("Enter the resizing factor (e.g., 1.5 for 1.5x): "))

    # Resize font and save the modified document
    resize_font(docx_path, resize_factor)
