
from docx import Document
import os
import sys
sys.path.insert(0, './../..')
from amplifydocx import resize_font


# Create a test Word document
def create_test_docx(file_path):
    doc = Document()
    doc.add_heading('Test Document', level=1)
    doc.add_paragraph('This is a sample paragraph in the test document.')

    # Save the document
    doc.save(file_path)

# Function to print the size of letters in the document
def print_letter_size(file_path):
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            print(f"Text: {run.text}, Size: {run.font.size}")

if __name__ == "__main__":
    # File paths
    test_docx_path = "test_document.docx"
    amplified_docx_path = "amplified_document.docx"

    # Create a test document
    create_test_docx(test_docx_path)

    # Amplify the document (hypothetical function)
    resize_font(test_docx_path, 2)  # Assuming the amplification factor is 2

    # Print the size of letters in the amplified document
    print("Size of letters in the amplified document:")
    print_letter_size(test_docx_path)

    # Clean up: Remove the test and amplified documents
    #os.remove(test_docx_path)
    #os.remove(amplified_docx_path)
